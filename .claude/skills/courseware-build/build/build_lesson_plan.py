#!/usr/bin/env python3
"""Generate the Responsible Generative AI Basics Lesson Plan (LP) DOCX
in the Tertiary WSQ house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 2-day schedule table (Day 1 9:30am-6:30pm full teaching day;
Day 2 9:30am-4:00pm teaching then the 4:00-6:00pm Assessment block).
The Daily Schedule table's "Slides" column is read from slide_map.json,
written by build_slides.py, so the LP always cites the actual deck page —
ALWAYS run build_slides.py before this script.
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT=DOMAIN1+DOMAIN2+DOMAIN3
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

with open(os.path.join(HERE,"slide_map.json")) as f:
    SLIDE_MAP=json.load(f)

def act_title(n): return next(a["title"] for a in ACT if a["num"]==n)
def slides_for(*keys):
    pages=[SLIDE_MAP[k] for k in keys if SLIDE_MAP.get(k)]
    if not pages: return "—"
    lo,hi=min(pages),max(pages)
    return f"Slide {lo}" if lo==hi else f"Slides {lo}–{hi}"

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"; ADMIN_FILL="F3F5F8"

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, text, slide_keys)  kind: admin/topic/activity/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","9:45",15,"admin","Digital Attendance (AM), welcome, trainer and learner introductions, ground rules, "
                              "Skills Framework and course outline",[]),
    ("9:45","10:20",35,"topic","LU1 Topic 1: Ethical considerations and potential risks of generative AI interaction (K2) — "
                               "hidden costs, deepfakes, environmental and societal impact, the deconstruction of AI bias",
                               ["topic1_section"]),
    ("10:20","10:40",20,"activity",f"Activity 1: {act_title(1)}",["activity1"]),
    ("10:40","11:05",25,"topic","LU1 Topic 2: Ethical principles in AI (K4) — AI Verify's 11 principles, the GenAI risk "
                                "diagnostic matrix, prompt-injection threat model, Project Moonshot",[]),
    ("11:05","11:15",10,"break","Tea break",[]),
    ("11:15","11:45",30,"activity",f"Activity 2: {act_title(2)} and Activity 3: {act_title(3)}",["activity2","activity3"]),
    ("11:45","12:10",25,"topic","LU1 Topic 3: Apply ethical principles in decision-making related to AI (A1) — three pillars of "
                                "algorithmic risk, national AI policy, the EU AI Act compliance waterfall, human in the loop",[]),
    ("12:10","12:55",45,"lunch","Lunch break",[]),
    ("12:55","1:20",25,"activity",f"Activity 4: {act_title(4)}",["activity4"]),
    ("1:20","1:40",20,"topic","LU1 Topic 4: Exercise professional scepticism to exercise sound judgement on GenAI output (A2) — "
                              "automation bias and the vulnerability in the human filter",[]),
    ("1:40","2:10",30,"activity",f"Activity 5: {act_title(5)} and Activity 6: {act_title(6)}; LU1 recap",["activity5","activity6"]),
    ("2:10","2:20",10,"admin","Digital Attendance (PM)",[]),
    ("2:20","2:45",25,"topic","LU2 Topic 1: Data anonymisation and de-identification techniques (K1) — anonymisation vs "
                              "pseudonymisation, k-anonymity, choosing the right protection",["topic2_section"]),
    ("2:45","3:05",20,"activity",f"Activity 7: {act_title(7)}",["activity7"]),
    ("3:05","3:30",25,"topic","LU2 Topics 2–3: Apply privacy measures when handling data for AI applications (A3) and design "
                              "guidelines for ethical AI use of sensitive data (A4) — the three layers of AI data security, "
                              "the 7-step governance checklist",[]),
    ("3:30","3:40",10,"break","Tea break",[]),
    ("3:40","4:10",30,"activity",f"Activity 8: {act_title(8)} and Activity 9: {act_title(9)}; LU2 recap",["activity8","activity9"]),
    ("4:10","4:35",25,"topic","LU3 Topics 1–2: Responsible AI principles and best practices, and comparing AI systems on IP, "
                              "data privacy and environmental impact (K3, A5) — EU AI Act risk taxonomy, ISO/IEC 42005, "
                              "differential privacy, Red AI vs Green AI",["topic3_section"]),
    ("4:35","5:05",30,"activity",f"Activities 10–12: {act_title(10)}, {act_title(11)} and {act_title(12)}",
                                 ["activity10","activity11","activity12"]),
    ("5:05","5:20",15,"topic","LU3 Topic 3: Compare ethical issues in AI applications (A6) — the anatomy of algorithmic "
                              "discrimination, accuracy vs interpretability, LIME vs SHAP",[]),
    ("5:20","5:45",25,"activity",f"Activities 13–14: {act_title(13)} and {act_title(14)}; LU3 recap",
                                 ["activity13","activity14"]),
    ("5:45","5:55",10,"recap","Course summary, what you achieved, Q&A and Courseware & Assessment on the LMS",[]),
    ("5:55","6:05",10,"assess","Briefing for Assessment and Digital Attendance (Assessment)",["briefing_assessment"]),
    ("6:05","6:35",30,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 30 minutes, open book (K1–K4)",
                               ["assessment_front"]),
    ("6:35","7:05",30,"assess","Case Study (CS) — 30 minutes, open book (A1–A6). TRAQOM survey and course feedback",
                               ["assessment_end"]),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.DOC_VERSION,
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,list(C.VERSION_HISTORY))
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("TSC Reference",f"{C.TSC_TITLE} ({C.TSC_CODE}) · {C.TSC_LEVEL}"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","1 day · 9:30am–7:05pm  (7.67 instructional hours + 1 hour of assessment)"),
      ("Daily Timing","45-minute lunch; two 10-minute tea breaks counted within training time"),
      ("Mode","Instructor-led, with case-study and workshop activities per Learning Unit"),
      ("Instructional Methods","Interactive presentation, discussions, case studies, peer teaching / peer practice"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "The assessment block runs at the end of the training day, from 6:05pm to 7:05pm.",
          "The Written Assessment covers the knowledge statements K1–K4; the Case Study covers the ability "
          "statements A1–A6. A candidate must be Competent in every K and A.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":ADMIN_FILL,"recap":ADMIN_FILL,"activity":None}

H("Daily Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=4); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity","Slides"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text,skeys in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[3],slides_for(*skeys),size=9.5,fill=fill,align=WD_ALIGN_PARAGRAPH.CENTER)
        if kind!="lunch": training+=mins
    for row in tbl.rows:
        row.cells[0].width=Inches(1.0); row.cells[1].width=Inches(0.8)
        row.cells[2].width=Inches(4.2); row.cells[3].width=Inches(1.1)
    assess_mins=sum(m for _s,_e,m,k,_t,_x in rows if k=="assess")
    instr_mins=training-assess_mins        # `training` already excludes lunch
    p=doc.add_paragraph(); r=p.add_run(
        f"Instructional time: {instr_mins} minutes ({instr_mins/60:.2f} hours), excluding the 45-minute "
        f"lunch.   Assessment block: {assess_mins} minutes (briefing + WA + CS).   "
        f"Total contact time: {training} minutes ({training/60:.2f} hours).")
    r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    p2=doc.add_paragraph(); r2=p2.add_run(
        "The assessment itself is 1 hour — Written Assessment (SAQ) 30 minutes plus Case Study 30 minutes "
        "— in line with Assessment Plan v2.0 (0.5 hrs per instrument, 1 hr total), preceded by a 10-minute "
        "assessment briefing. Digital attendance is taken three times — AM, PM and Assessment — as "
        "required for WSQ-funded courses.")
    r2.italic=True; r2.font.size=Pt(9.5); r2.font.color.rgb=GREY

H("Topic-by-Topic Breakdown",1)
for tp in C.TOPICS:
    topic_key=f"topic{tp['num']}_section"
    H(f"{tp['code']} — {tp['title']} · {slides_for(topic_key)}",2)
    doc.add_paragraph(tp["subtitle"])
    p=doc.add_paragraph("Key concepts:"); p.runs[0].bold=True
    for c in tp["concepts"]:
        doc.add_paragraph(c,style="List Bullet")
    for a in [x for x in ACT if x["topic"]==tp["num"]]:
        act_key=f"activity{a['num']}"
        H(f"Activity {a['num']} — {a['title']} · {slides_for(act_key)}",3)
        doc.add_paragraph(f"Objective: {a['objective']}.")
        doc.add_paragraph(f"Scenario: {a['desc']}")
        doc.add_paragraph(f"Duration: {a['duration']}.")

H("Resources Required",1)
for r in ["Projector/screen and PA system for the trainer's slide deck.",
          "Whiteboard/flip chart and markers for group activity work.",
          "Printed or digital copies of the Learner Guide for every learner.",
          "Internet access and a laptop per learner — all 14 in-class activities are browser-based web apps.",
          "A free Google Gemini API key (aistudio.google.com/api-keys) for Activities 6, 7, 13 and 14.",
          "Mobile phones (learners' own) for the SSG digital-attendance QR scans."]:
    doc.add_paragraph(r,style="List Bullet")

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Grading: Competent / Not Yet Competent.",
          "A minimum of 75% attendance (SSG Digital Attendance record) is required for funding eligibility.",
          "Learners submit their assessment answers and complete the TRAQOM survey on the LMS at "
          "https://lms-tms.tertiaryinfotech.com/."]:
    doc.add_paragraph(a,style="List Bullet")

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
