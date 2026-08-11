#!/usr/bin/env python3
"""Generate the Responsible Generative AI Basics Learner
Guide as a DOCX (courseware/LG-*.docx) — DOCX + PDF only, no Markdown mirror
kept in the repo (wsq-learner-guide HARD RULE 1).

House format: cover page, Document Version Control Record, auto TOC, Arial
11pt body, one section per Learning Unit with concepts, then one subsection
per activity (Objective · Scenario · Step-by-step · Debrief), a Quick
Reference table, Support section and the assessment flow. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned
with the slide deck and Lesson Plan.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT=DOMAIN1+DOMAIN2+DOMAIN3
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")
LAB_BASE="https://tertiarycourses.github.io/TGS-2025060472-Responsible-Generative-AI-Basics/"

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)

doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.DOC_VERSION,
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,list(C.VERSION_HISTORY))
prodoc.add_toc(doc)

def h3(text,color=BRAND):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=color
    return p

doc.add_heading("How to Use This Guide",level=1)
doc.add_paragraph(
    f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
    f"It supports the Skills Framework Technical Skill & Competency “{C.TSC_TITLE}” "
    f"({C.TSC_CODE}) across 3 Learning Units and 14 hands-on in-class activities.")
doc.add_paragraph(
    "Use this guide alongside the course slides during class, and again during the open-book "
    "assessment. Each Learning Unit section below lists the key concepts, followed by every "
    "activity you will complete in class with its scenario, its full step-by-step instructions "
    "and a debrief prompt to check your own work. The slides show only the high-level workflow "
    "of each activity — the detailed steps you actually follow are here in this guide.")
doc.add_paragraph(
    "All 14 activities are browser-based web apps. They run in any modern browser on your own "
    "laptop, and nothing needs to be installed. Open the activity hub at:")
p=doc.add_paragraph(); r=p.add_run(LAB_BASE); r.bold=True; r.font.color.rgb=BRAND
doc.add_paragraph("Before you start, you will need:")
for b in ["A laptop with a modern browser (Chrome, Edge, Firefox or Safari) and internet access.",
          "A free Google Gemini API key from https://aistudio.google.com/api-keys — needed for "
          "Activity 6 and Activity 7, and optional for Activities 13 and 14. The key is typed into "
          "the page, is never stored, and is only used from your own browser.",
          "A notebook or digital document to capture your activity results and reflections.",
          "The course slides (downloaded from the LMS) for reference during activities and the assessment."]:
    doc.add_paragraph(b,style="List Bullet")

doc.add_heading("Course Learning Outcomes",level=1)
for lo in C.LEARNING_OUTCOMES:
    doc.add_paragraph(lo,style="List Bullet")

doc.add_heading("Skills Framework Reference",level=1)
for line in [f"TSC Title: {C.TSC_TITLE}", f"TSC Code: {C.TSC_CODE}"]:
    doc.add_paragraph(line,style="List Bullet")
doc.add_paragraph(
    "You are assessed against the knowledge and ability statements below. The Written Assessment "
    "covers the knowledge statements (K1-K4); the Case Study covers the ability statements (A1-A6). "
    "You must be assessed Competent in every one of them.")
ktbl=doc.add_table(rows=0,cols=3); ktbl.style="Table Grid"
kh=ktbl.add_row().cells
for i,htext in enumerate(["Code","Statement","Assessed by"]):
    kh[i].text=""; r=kh[i].paragraphs[0].add_run(htext); r.bold=True; r.font.size=Pt(9.5)
    prodoc._shade_cell(kh[i],"1F6FEB"); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for code,stmt in list(C.KNOWLEDGE)+list(C.ABILITIES):
    cells=ktbl.add_row().cells
    cells[0].text=""; rr=cells[0].paragraphs[0].add_run(code); rr.bold=True; rr.font.size=Pt(9.5)
    cells[1].text=""; cells[1].paragraphs[0].add_run(stmt).font.size=Pt(9.5)
    cells[2].text=""; cells[2].paragraphs[0].add_run(
        "Written Assessment (SAQ)" if code.startswith("K") else "Case Study").font.size=Pt(9.5)

for t in C.TOPICS:
    doc.add_heading(f"{t['code']} — {t['title']}",level=1)
    doc.add_paragraph(t["subtitle"])
    h3("Key concepts")
    for c in t["concepts"]:
        doc.add_paragraph(c,style="List Bullet")
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        doc.add_heading(f"Activity {a['num']} — {a['title']}",level=2)
        doc.add_paragraph(f"Objective: {a['objective']}.")
        pl=doc.add_paragraph(); rl=pl.add_run("Open: "); rl.bold=True; rl.font.color.rgb=BRAND
        rl2=pl.add_run(LAB_BASE+a["path"].split("labs/",1)[1]); rl2.font.size=Pt(10)
        if a.get("apikey"):
            pk=doc.add_paragraph(); rk=pk.add_run("Requires: "); rk.bold=True; rk.font.color.rgb=BRAND
            pk.add_run(f"a {a['apikey']} API key — get a free key at https://aistudio.google.com/api-keys").font.size=Pt(10)
        h3("Scenario")
        doc.add_paragraph(a["desc"])
        h3("You'll produce")
        doc.add_paragraph(f"{a['build']}   (Duration: {a['duration']}.)")
        # Activity screenshot — what the learner should see when the app opens.
        shot=os.path.join(ASSETS,"labs",f"{a['slug']}.png")
        if os.path.exists(shot):
            doc.add_picture(shot,width=Inches(5.1))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rc=cap.add_run(f"Activity {a['num']} — {a['title']}")
            rc.italic=True; rc.font.size=Pt(9); rc.font.color.rgb=GREY
        h3("Step-by-step")
        # manual numbering so every activity restarts at step 1 (Word's
        # List Number style would otherwise run 1..71 across all activities)
        for i,(instr,_cmd) in enumerate(a["steps"],1):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Pt(18)
            r=p.add_run(f"{i}.  "); r.bold=True; r.font.color.rgb=BRAND
            p.add_run(instr)
        # "Test it" verification box — a real bordered, shaded single-cell table so
        # the learner can see at a glance how to self-check the activity.
        tb=doc.add_table(rows=1,cols=1); tb.style="Table Grid"
        cell=tb.rows[0].cells[0]
        prodoc._shade_cell(cell,"E8F7EE")
        cell.text=""
        ph=cell.paragraphs[0]; rh=ph.add_run("✅  Test it")
        rh.bold=True; rh.font.size=Pt(11); rh.font.color.rgb=RGBColor(0x12,0x7A,0x3E)
        pb=cell.add_paragraph(); pb.add_run(a["test"]).font.size=Pt(10.5)
        doc.add_paragraph("")

doc.add_heading("Quick Reference — Activities by Learning Unit",level=1)
tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"
hdr=tbl.add_row().cells
for i,htext in enumerate(["Learning Unit","Activity","Duration"]):
    hdr[i].text=""; r=hdr[i].paragraphs[0].add_run(htext); r.bold=True; r.font.size=Pt(9.5)
    prodoc._shade_cell(hdr[i],"1F6FEB")
    r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for t in C.TOPICS:
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        cells=tbl.add_row().cells
        cells[0].text=""; cells[0].paragraphs[0].add_run(t["code"]).font.size=Pt(9.5)
        cells[1].text=""; cells[1].paragraphs[0].add_run(f"{a['num']}. {a['title']}").font.size=Pt(9.5)
        cells[2].text=""; cells[2].paragraphs[0].add_run(a["duration"]).font.size=Pt(9.5)

doc.add_heading("Assessment",level=1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — this Learner Guide, the course slides and approved materials only.",
          "Grading: Competent / Not Yet Competent.",C.ASSESSMENT["note"]]:
    doc.add_paragraph(a,style="List Bullet")

doc.add_heading("Assessment Flow",level=1)
for i,step in enumerate(["TRAQOM — scan the TRAQOM QR code on the LMS and complete the survey.",
                          "Assessment Digital Attendance.",
                          "Assessment (Written Assessment + Case Study).",
                          "Submit the assessment answers on the LMS.",
                          "Sign the Assessment Summary Record."],1):
    doc.add_paragraph(f"{i}. {step}")
doc.add_paragraph("Courseware and the assessment are on the LMS: https://lms-tms.tertiaryinfotech.com/")

doc.add_heading("Glossary",level=1)
for term,defn in [
    ("Anonymisation","Irreversibly severing the link between data and an individual, so that no key exists to restore identity. Anonymised data is no longer personal data."),
    ("Pseudonymisation","Replacing identifiers with a token while retaining a key that can restore identity. Pseudonymised data REMAINS personal data under the PDPA and GDPR."),
    ("De-identification","Removing direct identifiers from a record. On its own it leaves a high re-identification risk through quasi-identifiers."),
    ("Quasi-identifier","A field that is not identifying alone (age, postcode, gender) but can identify an individual when combined with others."),
    ("K-anonymity","A guarantee that any individual cannot be distinguished from at least k-1 others, achieved by generalising or suppressing quasi-identifiers."),
    ("Differential privacy","A mathematical guarantee that an output does not reveal whether any individual's record was in the dataset, achieved by adding calibrated noise."),
    ("Epsilon (ε)","The differential-privacy budget. A lower epsilon means stronger privacy and a noisier, less useful answer."),
    ("DP-SGD","Differentially Private Stochastic Gradient Descent — training that clips each example's gradient and adds noise so no single record can be recovered."),
    ("Homomorphic encryption","Encryption that allows computation on data without decrypting it, protecting data during processing rather than only at rest."),
    ("Format-preserving encryption","Encryption that keeps a value's original format so downstream systems still accept it, while removing its identifying power."),
    ("Automation bias","The tendency to over-trust automated output and apply less scrutiny than to a human colleague's work."),
    ("Professional scepticism","The discipline of verifying AI output against source evidence rather than against the confidence of its presentation."),
    ("Hallucination","Confident, fluent AI output that is fabricated or wrong against the underlying facts or sources."),
    ("Prompt injection","An attack that embeds instructions in input to override an AI system's original instructions or extract its secrets."),
    ("AI Verify","Singapore's AI testing framework — 11 principles of trustworthy AI, split into process checks and technical tests."),
    ("Project Moonshot","The open-source toolkit that operationalises AI Verify testing — benchmarking, red-teaming and baseline testing for LLM applications."),
    ("EU AI Act","EU law regulating AI by risk tier — unacceptable, high, limited and minimal — with obligations scaled to the tier."),
    ("ISO/IEC 42005","A 2025 guidance standard for internally assessing an AI system's impact on individuals, groups and society across its lifecycle."),
    ("PDPA","Singapore's Personal Data Protection Act — the baseline obligation for handling personal data in Singapore organisations."),
    ("XAI","Explainable AI — techniques such as LIME and SHAP that make a black-box model's decision interpretable."),
    ("LIME","Local Interpretable Model-agnostic Explanations — fits a simple model locally around one prediction. Fast but stochastic."),
    ("SHAP","SHapley Additive exPlanations — assigns each feature its fair share of the outcome. Deterministic and additive."),
    ("Green AI / Red AI","Routing tasks to the smallest capable model to cut energy and cost, versus defaulting to the largest frontier model."),
    ("Human in the loop","Keeping a qualified, accountable human in a decision where errors are high-stakes or irreversible."),
    ("TRAQOM","The SSG post-course survey completed via the LMS, distinct from digital attendance."),
]:
    p=doc.add_paragraph(style="List Bullet")
    r=p.add_run(term+" — "); r.bold=True; p.add_run(defn)

doc.add_heading("Support",level=1)
for line in ["Email: enquiry@tertiaryinfotech.com","Tel: +65 6100 0613","Website: www.tertiarycourses.com.sg",
             "LMS: https://lms-tms.tertiaryinfotech.com/"]:
    doc.add_paragraph(line,style="List Bullet")

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
